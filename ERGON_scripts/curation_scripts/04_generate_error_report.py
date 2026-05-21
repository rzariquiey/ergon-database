from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
for attr in ('left_margin','right_margin','top_margin','bottom_margin'):
    setattr(section, attr, Cm(2.5))

# ── Helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def bold_run(para, text, size=11, color=None):
    r = para.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return r

def normal_run(para, text, size=10):
    r = para.add_run(text)
    r.font.size = Pt(size)
    return r

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
    return p

def add_table(doc, headers, rows, col_widths, header_bg='1F3564', stripe='EEF3FA'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        hdr_cells[i].width = Cm(w)
        set_cell_bg(hdr_cells[i], header_bg)
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for idx, row in enumerate(rows):
        cells = table.add_row().cells
        bg = stripe if idx % 2 == 0 else 'FFFFFF'
        for i, (val, w) in enumerate(zip(row, col_widths)):
            cells[i].width = Cm(w)
            set_cell_bg(cells[i], bg)
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
    return table

# ── PORTADA ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
bold_run(p, 'REPORTE DE INCIDENCIAS — BASE DE DATOS ERGON', 16, '1F3564')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
normal_run(p, f'Generado el {datetime.date.today().strftime("%d de %B de %Y")}', 10)

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run(
    'Este reporte documenta todas las fichas de la base de datos ERGON que presentan '
    'algún tipo de incidencia detectada durante el procesamiento automático: nombres de '
    'archivo no estándar (ya corregidos), errores en el nombre de la hoja de cálculo, '
    'inconsistencias entre el glottocode del archivo y su contenido, y fichas cuya lengua '
    'no pudo ser identificada con certeza a partir de la bibliografía disponible.'
)
r.font.size = Pt(10)
p.paragraph_format.space_after = Pt(12)

# ── RESUMEN EJECUTIVO ─────────────────────────────────────────────────────
heading(doc, '1. Resumen ejecutivo')

summary_data = [
    ('Nombres de archivo corregidos',   '3',  'Corrección automática aplicada. Verificar glottocode de nucl1302.'),
    ('Nombre de hoja incorrecto',       '8',  'La hoja fue copiada de otra ficha y no se actualizó su nombre.'),
    ('Contenido de otra lengua',        '1',  'El archivo dech1234 contiene datos de Dadibi (dadi1250).'),
    ('Inconsistencia de carpeta origen','1',  'mati1255.xlsx estaba en la carpeta mati1239 del ZIP original.'),
    ('Lengua no identificada',          '17', 'Metadatos marcados como VERIFICAR; requieren revisión manual.'),
    ('Posible error de glottocode',     '1',  'nucl1302: la fuente Hewitt 1990 no corresponde a Malagasy.'),
]
add_table(doc,
    ['Tipo de incidencia', 'N.º fichas', 'Descripción resumida'],
    summary_data,
    [7.5, 1.8, 7.0])

doc.add_paragraph()

# ── SECCIÓN 1: NOMBRES CORREGIDOS ─────────────────────────────────────────
heading(doc, '2. Nombres de archivo corregidos', 2)
p = doc.add_paragraph()
p.add_run(
    'Los siguientes archivos tenían nombres que incluían elementos distintos al glottocode '
    '(texto libre, nombre de investigador, etc.). El nombre ha sido corregido automáticamente '
    'al glottocode puro. Se recomienda verificar el caso de nucl1302.'
).font.size = Pt(10)

renamed_data = [
    ('nucl1302', 'Plantilla de Ergon - nucl1302.xlsx',
     'Renombrado a nucl1302.xlsx. ATENCIÓN: la fuente citada es Hewitt 1990, '
     'que corresponde a una lengua caucásica (georgiano o abjasio), no a Nuclear Malagasy. '
     'Verificar si el glottocode es correcto y si los datos son de ejemplo/plantilla.'),
    ('nina1238', 'Shiriana_language_ergativity.xlsx',
     'Renombrado a nina1238.xlsx. La ficha corresponde a Ninam/Shiriana '
     '(familia Yanomaman), confirmado por la fuente Gomez 1990.'),
    ('nort2745', 'nort2745 - Santiago Chau Principe.xlsx',
     'Renombrado a nort2745.xlsx. El sufijo "Santiago Chau Principe" era el nombre '
     'del investigador añadido al nombre del archivo.'),
]
add_table(doc,
    ['Glottocode', 'Nombre original', 'Observación'],
    renamed_data,
    [2.5, 5.0, 8.8])
doc.add_paragraph()

# ── SECCIÓN 2: NOMBRE DE HOJA INCORRECTO ──────────────────────────────────
heading(doc, '3. Nombre de hoja de cálculo incorrecto', 2)
p = doc.add_paragraph()
p.add_run(
    'Estas fichas fueron creadas copiando una plantilla existente sin actualizar el nombre '
    'de la hoja. El nombre de la hoja contiene el glottocode y/o el nombre de otra lengua '
    'distinta a la que debería corresponder el archivo. Los datos internos parecen ser '
    'correctos, pero el nombre de la hoja debe ser corregido manualmente.'
).font.size = Pt(10)

sheet_data = [
    ('caca1251', '"cent2127 - yupik"',
     'Kakataibo (Panoan). La hoja tiene el nombre de Central Alaskan Yupik. '
     'Contenido interno correcto (fuente: Zariquiey 2018 sobre Kakataibo).'),
    ('mati1255', '"cent2127 - yupik"',
     'Matis (Panoan). Además, el archivo estaba en la carpeta mati1239 del ZIP original '
     '(véase sección 5). Contenido interno correcto (fuente: Fleck 2010 sobre Matis).'),
    ('mats1244', '"cent2127 - yupik"',
     'Matsés (Panoan). Contenido interno correcto (fuente: Fleck 2010 sobre Matsés).'),
    ('yami1258', '"cent2127 - yupik"',
     'Yami (Austronesian). Contenido interno correcto (fuente: Neely 2019).'),
    ('dara1250', '"dadi1250_dadibi"',
     'Darai (Indo-European). La hoja tiene el nombre de Dadibi (dadi1250, Chimbu-Wahgi). '
     'Contenido interno correcto (fuente: Dhakal 2012 sobre Darai).'),
    ('darl1243', '"dadi1250_dadibi"',
     'Lengua de la cuenca del río Darling, posiblemente Paakantyi (Pama-Nyungan). '
     'La hoja tiene el nombre de Dadibi. Fuente: Hercus 1982.'),
    ('darm1243', '"dadi1250_dadibi"',
     'Darma (Trans-Himalayan). La hoja tiene el nombre de Dadibi. '
     'Contenido interno correcto (fuente: Willis 2019).'),
    ('dech1234', '"dadi1250_dadibi"',
     'CASO CRÍTICO: el nombre de la hoja indica Dadibi (dadi1250) y el contenido interno '
     'parece corresponder a esa lengua (véase sección 4). El glottocode del archivo '
     '(dech1234) y su contenido real no coinciden.'),
]
add_table(doc,
    ['Glottocode (archivo)', 'Nombre de hoja', 'Lengua real y observación'],
    sheet_data,
    [2.5, 3.5, 10.3])
doc.add_paragraph()

# ── SECCIÓN 3: CONTENIDO DE OTRA LENGUA ───────────────────────────────────
heading(doc, '4. Posible contenido de otra lengua', 2)
p = doc.add_paragraph()
p.add_run(
    'El siguiente archivo presenta una inconsistencia grave entre su nombre (glottocode) '
    'y el contenido interno de la ficha.'
).font.size = Pt(10)

content_data = [
    ('dech1234', 'dadi1250 (Dadibi)',
     'El nombre del archivo es dech1234, pero el nombre de la hoja es "dadi1250_dadibi" '
     'y la fuente citada ("Yu 2007") aparece también en ning1281. Los datos internos '
     '(comentarios, estructura de rasgos) son consistentes con Dadibi (familia Chimbu-Wahgi, PNG), '
     'no con la lengua que correspondería a dech1234. '
     'Acción recomendada: determinar si la ficha debe renombrarse a dadi1250.xlsx '
     '(verificando que no exista ya) o si los datos deben ser reemplazados por los de dech1234.'),
]
add_table(doc,
    ['Glottocode (archivo)', 'Lengua del contenido', 'Descripción del problema'],
    content_data,
    [2.5, 2.5, 11.3])
doc.add_paragraph()

# ── SECCIÓN 4: INCONSISTENCIA CARPETA ORIGEN ──────────────────────────────
heading(doc, '5. Inconsistencia de carpeta de origen', 2)
p = doc.add_paragraph()
p.add_run(
    'El siguiente archivo estaba ubicado en una subcarpeta del ZIP original cuyo nombre '
    '(glottocode de la carpeta) no coincide con el nombre del archivo.'
).font.size = Pt(10)

folder_data = [
    ('mati1255', 'mati1239',
     'El archivo mati1255.xlsx (Matis, Panoan; fuente: Fleck 2010) estaba alojado '
     'dentro de la carpeta "mati1239/" en el ZIP original. El glottocode mati1239 '
     'es distinto de mati1255. Los datos del archivo corresponden a Matis (mati1255), '
     'por lo que probablemente la carpeta tiene un nombre incorrecto en el repositorio de origen.'),
]
add_table(doc,
    ['Glottocode (archivo)', 'Glottocode (carpeta ZIP)', 'Descripción'],
    folder_data,
    [2.5, 2.8, 11.0])
doc.add_paragraph()

# ── SECCIÓN 5: LENGUA NO IDENTIFICADA ────────────────────────────────────
heading(doc, '6. Lengua no identificada — requiere verificación manual', 2)
p = doc.add_paragraph()
p.add_run(
    'Para los siguientes glottocodes no fue posible identificar con certeza el nombre '
    'de la lengua ni su familia a partir de la bibliografía disponible en la ficha. '
    'Las columnas Language y Family están marcadas como "VERIFICAR" en estas fichas. '
    'Se indica la fuente bibliográfica disponible como pista para la identificación.'
).font.size = Pt(10)

unid_data = [
    ('aben1249', 'Nitsch 2009',           'Lengua no identificada'),
    ('ahin1234', 'Santiago 2013',         'Lengua no identificada'),
    ('aime1238', 'Aiton 2016',            'Lengua no identificada'),
    ('akun1241', 'Aragon 2014',           'Lengua no identificada'),
    ('amam1246', 'Aki & Pennington 2013', 'Lengua no identificada'),
    ('ango1257', 'Casilimas Rojas 1995',  'Lengua no identificada; posiblemente lengua indígena de Colombia/Venezuela'),
    ('bala1316', 'Ozanne-Rivierre 1998',  'Probablemente lengua kanak de Nueva Caledonia (Austronesian), pero la variedad específica no fue determinada'),
    ('bamb1270', 'Campbell 1989',         'Lengua no identificada'),
    ('bara1357', 'Dhakal 2014',           'Probablemente lengua de Nepal; fuente apunta a autor nepalés'),
    ('bayb1234', 'Rubino 2005',           'Probablemente lengua austronesia de Filipinas'),
    ('bilb1241', 'Dempwolff 1909',        'Probablemente lengua austronesia; Dempwolff trabajó en Oceanía y África Oriental alemana'),
    ('boun1245', 'Merlan & Rumsey 1991',  'Probablemente lengua australiana o papuana'),
    ('cuti1242', 'Dos Anjos 2011: 283',   'Probablemente lengua sudamericana'),
    ('dech1234', 'Yu 2007 (ver sección 4)','El contenido parece corresponder a Dadibi (dadi1250); glottocode del archivo sin identificar'),
    ('dhar1248', 'Terril 2002',           'Probablemente lengua australiana o del Pacífico; Terrill 2002 podría ser la gramática de Kokota (Salomón) o de una lengua australiana'),
    ('fran1266', 'Engel & Bartholomew 1987', 'Probablemente lengua mesoamericana (Engel y Bartholomew son lingüistas del SIL en México)'),
    ('kara1476', 'McKelson 2004',         'Lengua no identificada'),
]
add_table(doc,
    ['Glottocode', 'Fuente disponible', 'Observación'],
    unid_data,
    [2.2, 3.8, 10.3])
doc.add_paragraph()

# ── SECCIÓN 6: POSIBLE ERROR GLOTTOCODE ───────────────────────────────────
heading(doc, '7. Posible error de glottocode', 2)
p = doc.add_paragraph()
p.add_run(
    'El siguiente archivo presenta una discrepancia entre el glottocode asignado y la '
    'bibliografía citada internamente.'
).font.size = Pt(10)

gc_err_data = [
    ('nucl1302', 'Nuclear Malagasy (Austronesian)',
     'La fuente citada es Hewitt 1990, autor conocido por trabajar sobre lenguas del Cáucaso '
     '(georgiano, abjasio). Esta fuente no corresponde a Nuclear Malagasy. '
     'Posibilidades: (a) el archivo es una plantilla de ejemplo con datos ficticios; '
     '(b) el glottocode asignado es incorrecto y la lengua real es georgiana (kart1248) '
     'u otra lengua caucásica. Verificar el origen del archivo.'),
]
add_table(doc,
    ['Glottocode', 'Lengua asignada', 'Descripción del problema'],
    gc_err_data,
    [2.2, 3.5, 10.6])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('— Fin del reporte —').bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

out = '/sessions/compassionate-lucid-gauss/mnt/Ergon/Reporte_incidencias_ERGON.docx'
doc.save(out)
print(f'Guardado: {out}')
